from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx2pdf import convert

# Few constants which are common
# Addresses
TEMPE = "Tempe, AZ 85287"
POLY = "Mesa, AZ 85212"
PHX = "Phoenix, AZ 85004"
WEST = "Glendale, AZ 85306"
# Campus names
TC = "Tempe Campus"
PC = "Polytechnic Campus"
DPC = "Downtown Phoenix Campus"
WC = "West Valley Campus"


def main():
    template_file_path = "CoverLetterTemplate.docx"
    output_file_path = "D:/Resumes/Cover_letters/97694BR_Data_Management_Assistant.docx"
    pdf_file_path = "D:/Resumes/Cover_letters/Pradnya_Chaudhari_Cover_Letter.pdf"

    today = date.today()
    formatted_date = today.strftime('%d %B, %Y')

    variables = {
        "${DATE}": formatted_date,
        "${DEPARTMENT}": "University Design Institute",
        "${ADDRESS}": TC,       # Campus Name
        "${CITY}": TEMPE        # Address
    }
    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

    para1 = template_document.add_paragraph(
        "I am writing to express my enthusiasm for the Data Management Assistant position (97694BR) at the University Design Institute. As a highly motivated and detail-oriented student with a keen interest in data management and system optimization, I am excited about the opportunity to contribute to the strategic development and maintenance of Airtable systems within UDI.")
    para1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    para2 = template_document.add_paragraph(
        "My academic pursuits in Computer Science have equipped me with a solid foundation in data management, relational databases, and technological solutions. Additionally, experiences gained from working as Software Engineer at MasterCard have honed my skills in scripting, automating data workflows, and maintaining data integrity—skills that closely align with the needs outlined for this role.")
    para2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    para3 = template_document.add_paragraph(
        "The responsibilities described for the Data Management Assistant position resonate deeply with my skill set and aspirations. I am proficient in supporting the ongoing design and development of Airtable systems, scripting and optimizing data workflows, and ensuring the quality assurance of data and systems. My ability to engage with institute staff to understand their data and system needs, coupled with my eagerness to learn and adapt, positions me well to thrive in this role.")
    para3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    para4 = template_document.add_paragraph(
        "I am particularly drawn to UDI's commitment to advancing its internal systems and ensuring technology platforms align with strategic objectives. Enclosed is my resume, highlighting my academic achievements and experiences relevant to data management and systems optimization.")
    para4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Uncomment when needed

    # para5 = template_document.add_paragraph(
    #     "Enclosed is my resume, highlighting my academic achievements and relevant experiences. I am excited about the prospect of bringing my dedication, creativity, and attention to detail to the EdPlus team.")
    # para5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # para6 = template_document.add_paragraph(
    #     "I am particularly excited about the opportunity to leverage my skills in Java, Spring Boot, Python, Django, IDEs like STS, IntelliJ IDEA, PyCharm, Git, JavaScript, JQuery, CSS, HTML, Jenkins, Jira and Confluence to contribute to your computational projects. I am committed to delivering efficient, scalable, and user-centric solutions while collaborating effectively with your team.")
    # para6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    last_para = template_document.add_paragraph(
        "Thank you for considering my application. I am eager to further discuss how my qualifications align with the needs of the Data Management Assistant position. I am available at your convenience for an interview and excited about the opportunity to contribute to the innovative work at the University Design Institute.")
    last_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Uncomment the following lines when you have email ID & contact number mentioned in cover letter

    # text = last_para.text
    # start = text.find("(xxx) xxx-xxxx")
    # prefix = text[:start]
    # suffix = text[start + len("(xxx) xxx-xxxx or xxxxxx@xxx.xxx"):]
    # last_para.text = prefix
    # run1 = last_para.add_run("(xxx) xxx-xxxx")
    # run1.font.color.theme_color = MSO_THEME_COLOR.HYPERLINK
    # last_para.add_run(" or ")
    # run2 = last_para.add_run("xxxxxx@xxx.xxx")
    # run2.font.color.theme_color = MSO_THEME_COLOR.HYPERLINK
    # run2.font.underline = True
    # last_para.add_run(suffix)

    salutation = template_document.add_paragraph("Warm regards, \nPradnya Chaudhari")

    template_document.save(output_file_path)
    convert(output_file_path, pdf_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()
